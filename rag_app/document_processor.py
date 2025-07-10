"""
Document Processor for RAG Application

This module provides document processing capabilities including file reading,
text extraction, chunking, and preprocessing for various document types.
"""

import os
import uuid
import re
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime
import hashlib

# File type imports
import pypdf
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter, TokenTextSplitter

from .vector_store import Document


class DocumentProcessor:
    """
    A processor for handling various document types and preparing them for RAG.
    """
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'docx',
        '.html': 'html',
        '.htm': 'html',
        '.json': 'json',
        '.csv': 'csv',
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.h': 'c',
        '.hpp': 'cpp'
    }
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        chunking_strategy: str = "recursive",
        max_tokens_per_chunk: Optional[int] = None
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Number of characters to overlap between chunks
            chunking_strategy: Strategy to use for chunking ('recursive', 'token', 'sentence')
            max_tokens_per_chunk: Maximum tokens per chunk (if using token-based chunking)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.chunking_strategy = chunking_strategy
        self.max_tokens_per_chunk = max_tokens_per_chunk
        
        # Initialize text splitter
        self._initialize_text_splitter()
    
    def _initialize_text_splitter(self):
        """Initialize the text splitter based on the chosen strategy."""
        if self.chunking_strategy == "recursive":
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        elif self.chunking_strategy == "token":
            if self.max_tokens_per_chunk is None:
                raise ValueError("max_tokens_per_chunk must be specified for token chunking")
            self.text_splitter = TokenTextSplitter(
                chunk_size=self.max_tokens_per_chunk,
                chunk_overlap=self.chunk_overlap
            )
        else:
            raise ValueError(f"Unsupported chunking strategy: {self.chunking_strategy}")
    
    def process_file(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Process a single file and return a list of document chunks.
        
        Args:
            file_path: Path to the file to process
            metadata: Optional metadata to attach to documents
            
        Returns:
            List of Document objects
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract file information
        file_info = self._extract_file_info(file_path)
        
        # Read file content
        content = self._read_file(file_path)
        
        # Preprocess content
        content = self._preprocess_text(content)
        
        # Create base metadata
        base_metadata = {
            "source": file_path,
            "file_type": file_info["extension"],
            "file_size": file_info["size"],
            "created_at": datetime.now().isoformat(),
            "content_hash": self._calculate_content_hash(content)
        }
        
        # Add custom metadata
        if metadata:
            base_metadata.update(metadata)
        
        # Chunk the content
        chunks = self._chunk_text(content)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc_id = f"{file_info['name']}_{i}_{uuid.uuid4().hex[:8]}"
            
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "chunk_id": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk)
            })
            
            document = Document(
                id=doc_id,
                content=chunk,
                metadata=chunk_metadata
            )
            documents.append(document)
        
        return documents
    
    def process_directory(
        self, 
        directory_path: str, 
        recursive: bool = True,
        file_patterns: Optional[List[str]] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Process all supported files in a directory.
        
        Args:
            directory_path: Path to directory to process
            recursive: Whether to process subdirectories
            file_patterns: Optional list of file patterns to match
            metadata: Optional metadata to attach to documents
            
        Returns:
            List of Document objects
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        documents = []
        
        # Get all files
        files = self._get_files_in_directory(directory_path, recursive, file_patterns)
        
        for file_path in files:
            try:
                file_documents = self.process_file(file_path, metadata)
                documents.extend(file_documents)
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")
                continue
        
        return documents
    
    def process_text(self, text: str, source: str = "text_input", metadata: Optional[Dict[str, Any]] = None) -> List[Document]:
        """
        Process raw text content.
        
        Args:
            text: Raw text content
            source: Source identifier for the text
            metadata: Optional metadata to attach to documents
            
        Returns:
            List of Document objects
        """
        # Preprocess content
        content = self._preprocess_text(text)
        
        # Create base metadata
        base_metadata = {
            "source": source,
            "file_type": "text",
            "created_at": datetime.now().isoformat(),
            "content_hash": self._calculate_content_hash(content)
        }
        
        # Add custom metadata
        if metadata:
            base_metadata.update(metadata)
        
        # Chunk the content
        chunks = self._chunk_text(content)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc_id = f"{source}_{i}_{uuid.uuid4().hex[:8]}"
            
            chunk_metadata = base_metadata.copy()
            chunk_metadata.update({
                "chunk_id": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk)
            })
            
            document = Document(
                id=doc_id,
                content=chunk,
                metadata=chunk_metadata
            )
            documents.append(document)
        
        return documents
    
    def _extract_file_info(self, file_path: str) -> Dict[str, Any]:
        """Extract basic information about a file."""
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "name": path.stem,
            "extension": path.suffix.lower(),
            "size": stat.st_size,
            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "created": datetime.fromtimestamp(stat.st_ctime).isoformat()
        }
    
    def _read_file(self, file_path: str) -> str:
        """Read content from a file based on its type."""
        extension = Path(file_path).suffix.lower()
        
        if extension == '.pdf':
            return self._read_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return self._read_docx(file_path)
        elif extension in ['.txt', '.md', '.html', '.htm', '.json', '.csv', '.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.hpp']:
            return self._read_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _read_pdf(self, file_path: str) -> str:
        """Read content from a PDF file."""
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
    
    def _read_docx(self, file_path: str) -> str:
        """Read content from a Word document."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _read_text_file(self, file_path: str) -> str:
        """Read content from a text file."""
        encodings = ['utf-8', 'latin-1', 'ascii', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file: {file_path}")
    
    def _preprocess_text(self, text: str) -> str:
        """Preprocess text content."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove very short lines (likely artifacts)
        lines = text.split('\n')
        cleaned_lines = [line.strip() for line in lines if len(line.strip()) > 3]
        
        # Remove duplicate consecutive lines
        deduplicated_lines = []
        prev_line = ""
        for line in cleaned_lines:
            if line != prev_line:
                deduplicated_lines.append(line)
                prev_line = line
        
        # Join lines back
        text = '\n'.join(deduplicated_lines)
        
        # Final cleanup
        text = text.strip()
        
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks."""
        if not text.strip():
            return []
        
        chunks = self.text_splitter.split_text(text)
        
        # Filter out very short chunks
        chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
        
        return chunks
    
    def _calculate_content_hash(self, content: str) -> str:
        """Calculate MD5 hash of content."""
        return hashlib.md5(content.encode()).hexdigest()
    
    def _get_files_in_directory(
        self, 
        directory_path: str, 
        recursive: bool = True,
        file_patterns: Optional[List[str]] = None
    ) -> List[str]:
        """Get all supported files in a directory."""
        files = []
        
        if recursive:
            for root, dirs, filenames in os.walk(directory_path):
                for filename in filenames:
                    file_path = os.path.join(root, filename)
                    if self._is_supported_file(file_path, file_patterns):
                        files.append(file_path)
        else:
            for filename in os.listdir(directory_path):
                file_path = os.path.join(directory_path, filename)
                if os.path.isfile(file_path) and self._is_supported_file(file_path, file_patterns):
                    files.append(file_path)
        
        return files
    
    def _is_supported_file(self, file_path: str, file_patterns: Optional[List[str]] = None) -> bool:
        """Check if a file is supported."""
        extension = Path(file_path).suffix.lower()
        
        # Check if extension is supported
        if extension not in self.SUPPORTED_EXTENSIONS:
            return False
        
        # Check file patterns if provided
        if file_patterns:
            filename = Path(file_path).name
            for pattern in file_patterns:
                if re.match(pattern, filename):
                    return True
            return False
        
        return True
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.SUPPORTED_EXTENSIONS.keys())
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get statistics about the document processor."""
        return {
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "chunking_strategy": self.chunking_strategy,
            "max_tokens_per_chunk": self.max_tokens_per_chunk,
            "supported_extensions": self.get_supported_extensions()
        } 